"""
Script to create example COA Word templates for docxtpl.

Generates two templates:
    1. ``coa_template.docx`` — a simple template with Jinja2 placeholders
       matching the predefined COA section keys.
    2. Can be used as a starting point for users to customise.

Run:
    python -m modules.create_template
"""

import os

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from modules.coa_structure import COA_SECTIONS, COA_FIELD_KEYS

TEMPLATE_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), "templates")
TEMPLATE_PATH = os.path.join(TEMPLATE_DIR, "coa_template.docx")


def create_template():
    """Create a COA template .docx with Jinja2 placeholders for docxtpl."""
    os.makedirs(TEMPLATE_DIR, exist_ok=True)

    doc = Document()

    # Default style
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(1.5)

    # Title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run("СЕРТИФИКАТ АНАЛИЗА")
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = "Times New Roman"

    # Subtitle
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("(Перевод на русский язык)")
    r.font.size = Pt(11)
    r.font.name = "Times New Roman"
    r.font.color.rgb = RGBColor(100, 100, 100)

    # Minimal metadata (optional)
    doc.add_paragraph()
    meta_fields = [
        ("Исходный файл:", "{{ original_filename }}"),
        ("Дата перевода:", "{{ translation_date }}"),
    ]
    for label, placeholder in meta_fields:
        p = doc.add_paragraph()
        lr = p.add_run(f"{label} ")
        lr.bold = True
        lr.font.size = Pt(9)
        lr.font.name = "Times New Roman"
        lr.font.color.rgb = RGBColor(80, 80, 80)
        vr = p.add_run(placeholder)
        vr.font.size = Pt(9)
        vr.font.name = "Times New Roman"
        vr.font.color.rgb = RGBColor(80, 80, 80)

    # Divider
    div = doc.add_paragraph()
    div.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dr = div.add_run("─" * 70)
    dr.font.size = Pt(7)
    dr.font.color.rgb = RGBColor(180, 180, 180)

    # Section placeholders
    for key, label, _desc, _is_table in COA_SECTIONS:
        heading = doc.add_paragraph()
        heading.paragraph_format.space_before = Pt(10)
        heading.paragraph_format.space_after = Pt(4)
        hr = heading.add_run(label)
        hr.bold = True
        hr.font.size = Pt(12)
        hr.font.name = "Times New Roman"

        content = doc.add_paragraph()
        cr = content.add_run("{{ " + key + " }}")
        cr.font.name = "Times New Roman"
        cr.font.size = Pt(11)

    # Bottom divider
    doc.add_paragraph()
    div2 = doc.add_paragraph()
    div2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dr2 = div2.add_run("─" * 70)
    dr2.font.size = Pt(7)
    dr2.font.color.rgb = RGBColor(180, 180, 180)

    doc.save(TEMPLATE_PATH)
    print(f"Template created at: {TEMPLATE_PATH}")
    print(f"Sections ({len(COA_SECTIONS)}):")
    for key, label, _, _ in COA_SECTIONS:
        print(f"  {{{{ {key} }}}}  →  {label}")


if __name__ == "__main__":
    create_template()
